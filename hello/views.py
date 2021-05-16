from django.shortcuts import render, redirect
from .models import *
from .forms import *
from django.views.generic.edit import UpdateView, CreateView, DeleteView
from django.views.generic.detail import DetailView
from django.views.generic import ListView
from django.urls import reverse
from django.http import FileResponse
from django.http import HttpResponse, JsonResponse, FileResponse
from django.views.decorators.http import require_POST
from django.template.loader import render_to_string
from django import forms
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackQueryHandler
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
import requests
import telebot
from telebot import types
import telegram
from django.contrib.auth import authenticate, login
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.views import PasswordChangeDoneView
from django.contrib.auth.password_validation import validate_password, password_changed
from django.contrib.auth.models import User
from docx import Document
import os
from playsound import playsound
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
from django.forms import modelformset_factory
class PasswordReset(PasswordChangeDoneView):
    success_url = '/folder/'





def login(request):

    
    username = request.user.username
    print(username)
    password = request.user.password
    print(password_changed('m3sh13280706django'))
    user = authenticate(request, username=username, password=password)
    
    if user is not None:
        return login(request)
    else:
        return redirect('folder')


@permission_required('hello.add_account')
@login_required
def folder(request):
    search_dir = '/home/xiidot/x/xiidot/heroku/examplebot/files/main'
    os.chdir(search_dir)
    files = filter(os.path.isfile, os.listdir(search_dir))
    files = [os.path.join(search_dir, f) for f in files] # add path to each file
    files.sort(key=lambda x: os.path.getmtime(x))
    os.remove(files[0])
    profiles = Profile.objects.all()
    prefix = []
    for p in profiles:
        prefix.append(p.prefix)
    card_number = []
    for c in profiles:
        card_number.append(c.card_number)
    table = []
    for i in profiles:
        a = Account.objects.filter(pseudonym=i.pseudonym)
        table.append(a)
    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    ids = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]
    context = {'prefix': prefix, 'card_number': card_number, 'profiles': profiles, 'ps': 'Все', 'allprofiles': profiles, 'path': 'Все', 'table': table, 'ids': ids, 'months': months, 'years': [2017, 2018, 2019, 2020]}
    return render(request, 'bot/folder.html', context)
@login_required
def sortfolder(request, ps):
    profiles = Profile.objects.all()
    if ps != 'Все':
        profile = Profile.objects.get(pseudonym=ps)
        profiles = [profile]
    allprofiles = Profile.objects.all()
    # -----
    prefix = []
    for p in profiles:
        prefix.append(p.prefix)
    card_number = []
    for c in profiles:
        card_number.append(c.card_number)
    table = []
    for i in profiles:
        a = Account.objects.filter(pseudonym=i.pseudonym)
        table.append(a)
    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    ids = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]
    #========
    context = {'prefix': prefix, 'card_number': card_number, 'profiles': profiles, 'ps': ps, 'allprofiles': allprofiles, 'path': ps, 'table': table, 'ids': ids, 'months': months, 'years': [2017, 2018, 2019, 2020]}
    
    
    return render(request, 'bot/folder.html', context)


def accounts(request, pse):
    accs = Account.objects.filter(pseudonym=pse)
    context = {'accounts': accs}
    return render(request, 'bot/account.html', context)

def allaccounts(request):
    obj = Account.objects.all().order_by('-published', 'pseudonym')
    context = {'accounts': obj, 'ps': 'Все', 'status':'Все', 'year':'Все', 'month':'1'}
    context['years'] = ['Все', '2016', '2017', '2018', '2019', '2020', '2021', '2022', '2023', '2024', '2025', '2026', '2027']
    context['months'] = ['Все', 'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    context['statuses'] = ['Все', 'В работе', 'В ожидании', 'Оплачен', 'Отклонена', 'Нет']
    context['profiles'] = Profile.objects.all().order_by('pseudonym')
    context['path'] = 'Все'
    return render(request, 'bot/allaccount.html', context)
@login_required
def sortallaccounts(request, ps, status, year, month):
    obj = Account.objects.all().order_by('-published', 'pseudonym')
    if ps != 'Все':
        obj = obj.filter(pseudonym=ps)
    if status != 'Все':
        obj = obj.filter(status__s=status)
    try:
        obj = obj.filter(year=int(year))
    except:
        fwrferwfrwe=0
    
    if month != '1':
        obj = obj.filter(month=int(month)-1)
    

    context = {'accounts': obj, 'ps': ps, 'status': status, 'year': year, 'month': month}
    context['years'] = ['Все', '2016', '2017', '2018', '2019', '2020', '2021', '2022', '2023', '2024', '2025', '2026', '2027']
    context['months'] = ['Все', 'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    context['statuses'] = ['Все', 'В работе', 'В ожидании', 'Оплачен', 'Отклонена', 'Нет']
    context['profiles'] = Profile.objects.all().order_by('pseudonym')
    context['path'] = ps
    return render(request, 'bot/allaccount.html', context)

@permission_required('hello.can_add_account')
def addadmin(request):
    if request.method == 'POST':
        f = Addadmin(request.POST)
        if f.is_valid():
            User.objects.create_user(username=f.cleaned_data['username'], password=f.cleaned_data['password'], email=f.cleaned_data['email'])
            return redirect('folder')
        else:
            context = {'form': f}
            return render(request, 'bot/addadmin.html', context)
    else:
        f = Addadmin
        context = {'form': f}
        return render(request, 'bot/addadmin.html', context)

            


def Sendmessage(request, ps, issent):
    if request.method == 'POST':
        bbf = SendmessageForm(request.POST)
        if bbf.is_valid():
            
            if bbf.cleaned_data['select'] == 'all':
                msg = bbf.cleaned_data['message']
                users = subscribersbot.objects.all()
                for u in users:
                    my_token = '1173519806:AAHXIC_9zPjTtZpvH4_OG_3s6Gt4lkI3jYw'
                    bot = telegram.Bot(token=my_token)
                    bot.sendMessage(chat_id=u.user_id, text='Сообщения от админстратора:\n'+msg)

            else:
                pseu = bbf.cleaned_data['select']
                msg = bbf.cleaned_data['message']
                login = Profile.objects.get(pseudonym=pseu).login
                user_id = subscribersbot.objects.get(login=login).user_id
                my_token = '1173519806:AAHXIC_9zPjTtZpvH4_OG_3s6Gt4lkI3jYw'


                bot = telegram.Bot(token=my_token)
                
                bot.sendMessage(chat_id=user_id, text='Сообщения от админстратора:\n'+msg)

                return redirect('sendmessage', permanent=True, ps=ps, issent='yes')
            
            
            return redirect('profiles', permanent=False)
            
        else:
            profiles = Profile.objects.all()
            try:
                current_profil = Profile.objects.get(pseudonym=ps).pk
            except:
                current_profil = None
            context = {'form': bbf, 'ps': ps, 'profiles': profiles, 'issent': issent, 'current': current_profil}
            return render(request, 'bot/sendmessage.html', context)

    else:
        bbf = SendmessageForm()
        profiles = Profile.objects.all()
        try:
                current_profil = Profile.objects.get(pseudonym=ps).pk
        except:
            current_profil = None
        context = {'form': bbf, 'ps': ps, 'profiles': profiles, 'issent': issent, 'current': current_profil}
        return render(request, 'bot/sendmessage.html', context)    


class AccCreateView(LoginRequiredMixin, CreateView, PermissionRequiredMixin):
    template_name = 'bot/createacc.html'
    form_class = AccForm
    success_url = '/detail/{id}'
    permanent = True
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        
        context['profiles'] = Profile.objects.all().order_by("pseudonym")
        

        
        
        context['card'] = card.objects.all()
        return context


class ProfDetailView(DetailView):
    model = Profile
    
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        obj = context['object']
        
        if card.objects.filter(pseudonym=obj.pseudonym):
            context['h1'] = 'Данные были успешно изменены'
            card.objects.get_or_create(card_number=obj.card_number, pseudonym=obj.pseudonym)
        else:
            context['h1'] = 'Новый пользователь успешно добавлен'
            card.objects.create(card_number=obj.card_number, pseudonym=obj.pseudonym)
        document = Document('files/Договор финал.docx')    
        for p in document.paragraphs:
            print(p.text)
            if '{Ф.И.О. партнера}' in p.text:
                print(p.text)
                p.text = p.text.replace('{Ф.И.О. партнера}', obj.name)
        
        document.save('files/{}.docx'.format(obj.pseudonym))
        os.system("unoconv -f pdf files/{}.docx".format(obj.pseudonym))

        now = Profile.objects.filter(pk=obj.pk).values()[0]
        c = 0
        try:
            pre = prof_pre_value.objects.filter(pk=obj.pk).values()[0]
            for i in pre:
                if i == 'id':
                    asddw = 0
                elif pre[i] != now[i]:
                    stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Профил', text=i + ':\n' + str(pre[i]) + ' >> ' + str(now[i]))
                
                c += 1
        except:
            stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Профил', text="Создан")

        try:
            pre = prof_pre_value.objects.get(pk=obj.pk)
            pre.login = obj.login
            pre.parol = obj.parol
            pre.pseudonym = obj.pseudonym
            pre.number = obj.number
            pre.code = obj.code
            pre.name = obj.name
            pre.prefix = obj.prefix
            pre.type_payment = obj.type_payment
            pre.valute_card = obj.valute_card
            pre.card_number = obj.card_number
            pre.date_card = obj.date_card
            pre.owner_card = obj.owner_card
            pre.save()
        except:
            prof_pre_value.objects.create(pk=obj.pk, login = obj.login, parol = obj.parol, pseudonym = obj.pseudonym, number = obj.number, code = obj.code, name = obj.name, prefix = obj.prefix, type_payment = obj.type_payment, valute_card = obj.valute_card, card_number = obj.card_number, date_card = obj.date_card, owner_card = obj.owner_card)

        return context

class AccDetailView(DetailView):
    model = Account
    
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        
        
        obj = context['object']
        
        now = Account.objects.filter(pk=obj.pk).values()[0]
        c = 0
        try:
            pre = acc_pre_value.objects.filter(pk=obj.pk).values()[0]
        
            for i in pre:

                if i == 'id':
                    feef = 0
                elif i == 'published':
                    wefw = 0
                elif i == 'status':
                    try:
                        if pre[i] != obj.status.s:
                            stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text=i + ':\n' + pre[i] + ' >> ' + obj.status.s)        
                    except:
                        if pre[i] != '':
                            stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text=i + ':\n' + pre[i] + ' >> ' + "-------")
                
                elif i == 'month':
                    
                    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
                    month = months[now['month_id']-1]
                    if pre[i] != month:         
                        stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text=i + ':\n' + pre[i] + ' >> ' + month)
                elif pre[i] != now[i]:
                    
                    stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text=i + ':\n' + str(pre[i]) + ' >> ' + str(now[i]))
                
                c += 1
            action_story.objects.create(username=self.request.user, pseudonym=obj.pseudonym, obj='Отчет', url='/update/{}'.format(obj.pk))
        except:
            stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text="Создан")

            
        # save currently fields as old , that next time will be pre
        try:
            
            pre = acc_pre_value.objects.get(pk=obj.pk)
            if pre.pseudonym != obj.pseudonym:            
                pre.pseudonym = obj.pseudonym

            if str(pre.month) != str(obj.month):
                pre.month = str(obj.month)

            if pre.year != obj.year:    
                pre.year = (obj.year)

            try:
                if pre.status != obj.status.s:
                    pre.status = obj.status.s

            except:
                pre.status = ''

            if pre.day_payment != obj.day_payment:
                pre.day_payment = obj.day_payment


            if pre.document != obj.document:    
                pre.document = obj.document

        
            if pre.summa != obj.summa:
                pre.summa = obj.summa

            
            if pre.card_number != obj.card_number:
                pre.card_number = obj.card_number


            pre.save()
            
        except:
            try:
                status = obj.status.s
            except:
                status = ''
            pre = acc_pre_value.objects.create(pk=obj.pk, pseudonym = obj.pseudonym, month = obj.month, year = obj.year, status = status, day_payment = obj.day_payment, document = obj.document, summa = obj.summa, card_number = obj.card_number)

        return context


class ProfCreateView(CreateView):
    template_name = 'bot/createprof.html'
    form_class = ProfForm
    success_url = '/new/{id}'
    permanent = True
    def get_context_data(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        con['profiles'] = Profile.objects.all()
        return con





class AccDeleteView(DeleteView):
    model = Account
    def get_success_url(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']



        url = '/account/{}/'.format(obj.pseudonym)
        return url

class ProfDeleteView(DeleteView):
    model = Profile
    
    def get_success_url(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']
        ps = obj.pseudonym
        login = obj.login


        url = '/afterdeleting/{}/{}/'.format(ps, login)
        return url
   
def afterdeleting(request, ps, login):
    for i in Account.objects.filter(pseudonym=ps):
        i.delete()
    try:
        s = subscribersbot.objects.get(login=login)
        user_id = s.user_id
        s.delete()
        ch = changing.objects.get(user_id=user_id)
        ch.delete()
        t = typing.objects.get(user_id=user_id)
        t.delete()
        requests.get('https://api.telegram.org/bot/sendMessage?chat_id={}&text=ваш профиль удален\nнажмите /start, чтобы войти снова'.format(user_id))
    except:
        fwrfwe = 0
    return redirect(Profiles)


class AccEditView(UpdateView):
    model = Account
    form_class = AccForm
    success_url = '/detail/{id}'
    def get_context_data(self, *args, **kwargs):
        
        con = super().get_context_data(*args, **kwargs)
        con['profiles'] = Profile.objects.all()
        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
        con['months'] = months
        obj = con['object']
        con['stories'] = stories.objects.filter(obj_id=obj.pk, obj='Отчет')
        con['card'] = card.objects.filter(pseudonym=obj.pseudonym)
        return con


class ProfEditView(LoginRequiredMixin, UpdateView):
    model = Profile
    form_class = ProfForm
    success_url = '/new/{id}'
    def get_context_data(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']
        con['stories'] = stories.objects.filter(obj_id=obj.pk, obj='Профил')
        
        return con

@login_required
def Profiles(request):
    profiles = Profile.objects.all()
    bot_users = []
    for i in subscribersbot.objects.all():
        bot_users.append(i.login)
    context = {"profiles": profiles, 'ps': 'Все', 'allprofiles': profiles, 'path': 'Все', 'by': 'pseudonym', 'bot_users': bot_users}
    
    return render(request, 'bot/profiles.html', context)



@login_required
def SortProfiles(request, ps, by):
    allprofiles = Profile.objects.all()
    if ps == 'Все':
        profiles = Profile.objects.all().order_by('pseudonym').order_by(by)
    else:
        profiles = Profile.objects.filter(pseudonym=ps)

    context = {"profiles": profiles, 'ps': ps, 'allprofiles': allprofiles, 'path': ps, 'by': by}
    return render(request, 'bot/profiles.html', context)    

#______bot:
def bot(request):
    
    
    return redirect(open_site)

def get_file(request, file):
    print(file)
    f = open('/home/xiidot/x/xiidot/heroku/examplebot/files/'+file, 'rb')
    return FileResponse(f)


def open_site(request):
    if request.method == 'POST':
        f = SecurityForm(request.POST)
        
        if f.is_valid():
            
            try:
                parol = security.objects.get(pk=1).parol
                if f.cleaned_data['parol'] == parol:
                    return redirect('folder')
                else:
                    return render(request, 'bot/security.html', {'form': f, 'message': 'неверно, пожалуйста, введите пароль еще раз'})
            except:
                
                security.objects.create(parol = f.cleaned_data['parol'])
                return redirect('folder')
        else:
            return render(request, 'bot/security.html', {'form': f, 'message':'Введите пароль'})
    else:
        f = SecurityForm()
        return render(request, 'bot/security.html', {'form': f, 'message':'Введите пароль'})


class ContractEditView(UpdateView):

    model = contract
    form_class = SendContract
    success_url = '/'
    
class SecEditView(UpdateView):
    model = security
    form_class = SecurityForm
    success_url = '/'


def ActionStory(request):
    obj = action_story.objects.all()
    context = {'actions': obj}
    return render(request, 'bot/actionstory.html', context)

@login_required
def sendfile(request, y, m, d, f):
    document='{}/{}/{}/{}'.format(y, m, d, f)
    print(document)
    a = Account.objects.filter(document='{}/{}/{}/{}'.format(y, m, d, f))
    return FileResponse(a[0].document)

@login_required
def contract(request, ps):
    obj = Profile.objects.get(pseudonym=ps)
    document = Document('files/Договор финал.docx')
    for p in document.paragraphs:
        if '{Ф.И.О. партнера}' in p.text:
            p.text = p.text.replace('{Ф.И.О. партнера}', obj.name)

    document.save('files/{}.docx'.format(obj.pseudonym))
    p = os.path.abspath('files/{}.docx'.format(obj.pseudonym))
    os.system("unoconv -f pdf {}".format(p))
    f = open('files/{}.pdf'.format(ps), 'rb')
    return FileResponse(f)

def audio_create(request, artist):
    FormSet = modelformset_factory(Audio, fields=('composition', 'artist', 'autor_music', 'autor_text', 'album', 'isrc', 'genre', 'copyright', 'related_rights', 'upc', 'release_date', 'territory', 'link'),
        labels = {
            'composition': '', 
            'artist': '', 
            'autor_music': '', 
            'autor_text': '', 
            'album': '', 
            'isrc': '', 
            'genre': '', 
            'copyright': '', 
            'related_rights': '', 
            'upc': '', 
            'release_date': '', 
            'territory': '', 
            'link': '',
        },
        widgets = {
            'composition': forms.TextInput(attrs={'class': 'composition'}),
            'artist': forms.TextInput(attrs={'class': 'artist'}),
            'autor_music': forms.TextInput(attrs={'class': 'autor_music'}),
            'autor_text': forms.TextInput(attrs={'class': 'autor_text'}),
            'album': forms.TextInput(attrs={'class': 'album'}),
            'isrc': forms.TextInput(attrs={'class': 'isrc'}),
            'genre': forms.TextInput(attrs={'class': 'genre'}),
            'copyright': forms.TextInput(attrs={'class': 'copyright'}),
            'related_rights': forms.TextInput(attrs={'class': 'related_rights'}),
            'upc': forms.TextInput(attrs={'class': 'upc'}),
            'release_date': forms.TextInput(attrs={'class': 'release_date'}),
            'territory': forms.TextInput(attrs={'class': 'territory'}),
            'link': forms.TextInput(attrs={'class': 'link'}),
        }

    )
    if request.method == 'POST':
        formset = FormSet(request.POST, queryset=Audio.objects.filter(pseudonym=artist))
        if formset.is_valid():
            formset.save()
            for form in formset:
                if form.instance.pseudonym == None and form.instance.artist != None:
                    form.instance.pseudonym = artist
                    form.instance.status = '+'
                    current_pk = form.instance.pk
                    try:
                        video_pk = max(list(Video.objects.all().values_list('pk')))[0]
                        if current_pk > video_pk:
                            do_nothing=True
                        else:
                            form.instance.pk = video_pk + 1
                    except:
                        dewdfwe=921
                    form.save()
                #print(form.cleaned_data['artist'])
                #for f in form:
                #    print(form.artist)
        return redirect(content, ps=artist, pk=17)

    else:
        formset = FormSet(queryset=Audio.objects.filter(pseudonym=artist))
        count = len(Audio.objects.filter(pseudonym=artist))
        context = {'form': formset, 'count': count+1}
        return render(request, 'bot/createaudio.html', context)



def video_create(request, artist):
    FormSet = modelformset_factory(Video, fields=('composition', 'type', 'artist', 'isrc', 'producer', 'operator', 'autor_script', 'painter', 'copyright', 'release_date', 'territory', 'link'),
        labels = {
            'composition': '', 
            'type': '', 
            'artist': '', 
            'isrc': '', 
            'producer': '', 
            'operator': '', 
            'autor_script': '', 
            'painter': '', 
            'copyright': '', 
            'release_date': '', 
            'territory': '', 
            'link': '',
        },
        widgets = {
            'composition': forms.TextInput(attrs={'class': 'composition'}), 
            'type': forms.TextInput(attrs={'class': 'type'}), 
            'artist': forms.TextInput(attrs={'class': 'artist'}), 
            'isrc': forms.TextInput(attrs={'class': 'isrc'}), 
            'producer': forms.TextInput(attrs={'class': 'producer'}), 
            'operator': forms.TextInput(attrs={'class': 'operator'}), 
            'autor_script': forms.TextInput(attrs={'class': 'autor_script'}), 
            'painter': forms.TextInput(attrs={'class': 'painter'}), 
            'copyright': forms.TextInput(attrs={'class': 'copyright'}), 
            'release_date': forms.TextInput(attrs={'class': 'release_date'}),
            'territory': forms.TextInput(attrs={'class': 'territory'}), 
            'link': forms.TextInput(attrs={'class': 'link'}),
        }

    )
    if request.method == 'POST':
        formset = FormSet(request.POST, queryset=Video.objects.filter(pseudonym=artist))
        if formset.is_valid():
            formset.save()
            for form in formset:
                if form.instance.pseudonym == None and form.instance.artist != None:
                    form.instance.pseudonym = artist
                    current_pk = form.instance.pk
                    audio_pk = max(list(Audio.objects.all().values_list('pk')))[0]
                    print(audio_pk)
                    if current_pk > audio_pk:
                        do_nothing=True
                    else:
                        form.instance.pk = audio_pk + 1
                    form.save()
        return redirect(content, ps=artist)

    else:
        formset = FormSet(queryset=Video.objects.filter(pseudonym=artist))
        count = len(Video.objects.filter(pseudonym=artist))
        context = {'form': formset, 'count': count+1}
        return render(request, 'bot/createvideo.html', context)



class AudioCreateView(LoginRequiredMixin, CreateView):
    template_name = 'bot/createaudio.html'
    form_class = AudioForm
    success_url = '/audio_detail/{id}'
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = []
        for i in Profile.objects.all().order_by('pseudonym'):
            ps.append(i.pseudonym)
        context['profiles'] = ps
        context['allprofiles'] = Profile.objects.all().order_by('pseudonym')
        return context

class AuidoEditView(LoginRequiredMixin, UpdateView):
    model = Audio
    form_class = AudioForm
    def get_success_url(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = context['object'].artist
        return '/content/{}'.format(ps)

class AudioDetailView(LoginRequiredMixin, DetailView):
    model = Audio
@login_required
def delete_audio(request, pk):
    audio = Audio.objects.get(pk=pk)
    artist = audio.artist
    if audio.status == '+':
        audio.status = '-'
    else:
        audio.status = '+'
    audio.save()
    return redirect(content, ps=artist, pk=17)


class VideoCreateView(LoginRequiredMixin, CreateView):
    template_name = 'bot/createvideo.html'
    form_class = VideoForm
    success_url = '/video_detail/{id}'
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = []
        for i in Profile.objects.all().order_by('pseudonym'):
            ps.append(i.pseudonym)
        context['profiles'] = ps
        context['allprofiles'] = Profile.objects.all().order_by('pseudonym')
        return context


class VideoEditView(LoginRequiredMixin, UpdateView):
    model = Video
    form_class = VideoForm
    def get_success_url(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = context['object'].artist
        return '/content/{}'.format(ps)


class VideoDetailView(LoginRequiredMixin, DetailView):
    model = Video
@login_required
def delete_video(request, pk):
    video = Video.objects.get(pk=pk)
    artist = video.artist
    video.delete()
    return redirect(content, ps=artist)
@login_required
def content(request, ps, pk):
    audios = Audio.objects.filter(pseudonym=ps)
    indexes = {}

    for a in audios:
        obj = Audio.objects.filter(composition=a.composition, artist=a.artist, autor_music=a.autor_music, autor_text=a.autor_text, album=a.album, isrc=a.isrc, genre=a.genre, copyright=a.copyright, related_rights=a.related_rights, upc=a.upc, release_date=a.release_date, territory=a.territory)
        obj_pks = [i.pk for i in obj]
        index = obj_pks.index(a.pk) + 1
        indexes[a.pk] = index

    videos = Video.objects.filter(pseudonym=ps)
    indexes_video = {}

    for v in videos:
        obj = Video.objects.filter(composition=v.composition, artist=v.artist, isrc=v.isrc, copyright=v.copyright, release_date=v.release_date, territory=v.territory, type=v.type, producer=v.producer, operator=v.operator, autor_script=v.autor_script, painter=v.painter)
        obj_pks = [i.pk for i in obj]
        index = obj_pks.index(v.pk) + 1
        indexes_video[v.pk] = index

    context = {'audios': audios, 'videos': videos, 'ps': ps, 'pk': pk, 'indexes': indexes, 'indexes_video': indexes_video}
    return render(request, 'bot/content.html', context)

@login_required
def play_sound(request):
    playsound('/home/xiidot/x/xiidot/python/thecleverprogrammer.com/alarm clock/audio.m4a')
    return 0

@login_required
def generation_file(request, artist, pk):
    file_path = '/home/xiidot/x/xiidot/heroku/examplebot/files/app/Приложение шаблон.docx'
    folder_path = '/home/xiidot/x/xiidot/heroku/examplebot/files/app/'
    
    list_audio = list(Audio.objects.filter(pseudonym=artist).values_list())
    list_video = list(Video.objects.filter(pseudonym=artist).values_list())
    
    doc = Document(file_path)
    
    n = 1

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    def modifyBorder(table):  # set border size
        tbl = table._tbl # get xml element in table
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '30')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'black')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
    for table in doc.tables:
        if n == 2:
            modifyBorder(table)
            index = 1
            for l in list_audio:
                if l[-1] == '-':
                    continue
                rows = table.add_row()
                cn = 0
                for c in rows.cells:
                    if cn == 0:
                        c.text = str(index)
                    else:
                        c.text = l[cn]
                    cn += 1
                    p = c.paragraphs[0]    
                    for run in p.runs:
                        run.font.size = Pt(10)
                index += 1


        elif n == 3:
            modifyBorder(table)
            index = 1
            for l in list_video:
                rows = table.add_row()
                cn = 0
                for c in rows.cells:
                    if cn == 0:
                        c.text = str(index)
                    else:
                        c.text = l[cn]
                    cn += 1
                    p = c.paragraphs[0]    
                    for run in p.runs:
                        run.font.size = Pt(10)
                index += 1
        
        
        else:
            print(artist)
            artist = artist.replace('_', ' ')
            obj = Profile.objects.get(pseudonym=artist)
            for r in table.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        if obj.name != None:
                            p.text = p.text.replace('{name}', obj.name)
                        else:
                            p.text = p.text.replace('{name}', '________')
                        if obj.prefix != None:
                            p.text = p.text.replace('{prefix}', obj.prefix)
                        else:
                            p.text = p.text.replace('{prefix}', '_____')
                        artist = artist.replace(' ', '_')
                        try:
                            index = len(os.listdir(folder_path+artist))+1
                        except:
                            index = 1
                        p.text = p.text.replace('{index}', str(index))
        
                        y, m, d = str(obj.published).split('-')
                        p.text = p.text.replace('{date}', '{}.{}.{}'.format(d, m, y))
                        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
                        p.text = p.text.replace('{year}', y)
                        p.text = p.text.replace('{month}', months[int(m)-1])
                        p.text = p.text.replace('{day}', d)
                        

        n += 1
    try:
        d = datetime.now()
        date_time = '__{}-{}-{}-{}-{}-{}'.format(str(d.year), str(d.month), str(d.day), str(d.hour), str(d.minute), str(d.second))
        artist = artist.replace(' ', '_')
        doc.save(folder_path+artist+'/'+artist+date_time+'.docx')
    except:
        artist = artist.replace(' ', '_')
        os.system('mkdir {}'.format(folder_path+artist))
        doc.save(folder_path+artist+'/'+artist+date_time+'.docx')
    f = open(folder_path+artist+'/'+artist+date_time+'.docx', 'rb')
    return FileResponse(f)

@login_required
def app_list(request, artist):
    artist = artist.replace(' ', '_')
    all_apps = os.listdir('/home/xiidot/x/xiidot/heroku/examplebot/files/app/{}'.format(artist))
    date_time = []
    for i in all_apps:
        print(i)
        folder, d_t_docx = i.split('__')
        d_t, docx = d_t_docx.split('.')
        year, month, day, hour, minute, second = d_t.split('-')
        date_time.append('{}.{}.{} {}:{}'.format(day, month, year, hour, minute))
    context = {'all_apps': all_apps, 'date_time': date_time}
    return render(request, 'bot/app_list.html', context)

@login_required
def open_app(request, app):
    folder, date = app.split('__')
    folder = folder.replace(' ', '_')
    file_path = '/home/xiidot/x/xiidot/heroku/examplebot/files/app/'
    f = open(file_path+folder+'/'+app, 'rb')
    return FileResponse(f)
